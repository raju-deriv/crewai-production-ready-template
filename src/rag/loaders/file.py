import structlog
import os
import hashlib
import mimetypes
from typing import Dict, List, Any, Optional, Union
from src.rag.loaders.base import DocumentLoader

logger = structlog.get_logger(__name__)

class FileLoader(DocumentLoader):
    """
    Loader for local files.
    
    This class implements the DocumentLoader interface for local files.
    It provides methods for loading documents from file paths.
    """
    
    def __init__(self):
        """
        Initialize the FileLoader.
        """
        # Initialize supported file types
        self.supported_extensions = {
            '.txt': self._load_text,
            '.md': self._load_text,
            '.html': self._load_text,
            '.htm': self._load_text,
            '.csv': self._load_text,
            '.json': self._load_text,
            '.xml': self._load_text,
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.doc': self._load_docx,
            '.png': self._load_image,
            '.jpg': self._load_image,
            '.jpeg': self._load_image,
            '.gif': self._load_image,
            '.bmp': self._load_image
        }
        
        logger.info(f"Initialized FileLoader with {len(self.supported_extensions)} supported file types")
    
    def load(self, source: str, **kwargs) -> Dict[str, Any]:
        """
        Load a document from a file path.
        
        Args:
            source: File path to load.
            **kwargs: Additional arguments specific to the file type.
        
        Returns:
            Dict[str, Any]: The loaded document.
        """
        try:
            # Validate file path
            if not os.path.exists(source):
                raise FileNotFoundError(f"File not found: {source}")
            
            if not os.path.isfile(source):
                raise ValueError(f"Not a file: {source}")
            
            # Get file extension
            _, ext = os.path.splitext(source)
            ext = ext.lower()
            
            # Check if file type is supported
            if ext not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {ext}")
            
            # Load file using appropriate method
            loader_method = self.supported_extensions[ext]
            document = loader_method(source, **kwargs)
            
            # Generate document ID
            doc_id = f"file_{hashlib.md5(source.encode()).hexdigest()}"
            
            # Add document ID
            document['id'] = doc_id
            
            logger.info(f"Loaded file: {source}")
            return document
        except Exception as e:
            logger.error(f"Error loading file: {source}", error=str(e))
            # Return empty document with error metadata
            return {
                'id': f"file_{hashlib.md5(source.encode()).hexdigest()}",
                'text': '',
                'metadata': {
                    'source': source,
                    'error': str(e),
                    'source_type': 'file'
                }
            }
    
    def load_batch(self, sources: List[str], **kwargs) -> List[Dict[str, Any]]:
        """
        Load multiple documents from file paths.
        
        Args:
            sources: List of file paths to load.
            **kwargs: Additional arguments passed to load().
        
        Returns:
            List[Dict[str, Any]]: List of loaded documents.
        """
        documents = []
        
        for source in sources:
            document = self.load(source, **kwargs)
            documents.append(document)
        
        logger.info(f"Loaded {len(documents)} files")
        return documents
    
    def supports(self, source_type: str) -> bool:
        """
        Check if the loader supports a source type.
        
        Args:
            source_type: Type of source.
        
        Returns:
            bool: True if the loader supports the source type, False otherwise.
        """
        if source_type.lower() == 'file':
            return True
        
        # Check if it's a file extension
        if source_type.startswith('.'):
            return source_type.lower() in self.supported_extensions
        
        # Check if it's a mime type
        for ext in self.supported_extensions:
            mime_type, _ = mimetypes.guess_type(f"file{ext}")
            if mime_type and mime_type.lower() == source_type.lower():
                return True
        
        return False
    
    def _load_text(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Load a text file.
        
        Args:
            file_path: Path to the text file.
            **kwargs: Additional arguments.
        
        Returns:
            Dict[str, Any]: The loaded document.
        """
        encoding = kwargs.get('encoding', 'utf-8')
        
        with open(file_path, 'r', encoding=encoding) as f:
            text = f.read()
        
        # Get file metadata
        stat = os.stat(file_path)
        
        return {
            'text': text,
            'metadata': {
                'source': file_path,
                'filename': os.path.basename(file_path),
                'extension': os.path.splitext(file_path)[1].lower(),
                'size': stat.st_size,
                'modified': stat.st_mtime,
                'source_type': 'file'
            }
        }
    
    def _load_pdf(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Load a PDF file.
        
        Args:
            file_path: Path to the PDF file.
            **kwargs: Additional arguments.
        
        Returns:
            Dict[str, Any]: The loaded document.
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text = ""
            
            # Extract text from each page
            for page in reader.pages:
                text += page.extract_text() + "\n\n"
            
            # Get file metadata
            stat = os.stat(file_path)
            
            return {
                'text': text,
                'metadata': {
                    'source': file_path,
                    'filename': os.path.basename(file_path),
                    'extension': '.pdf',
                    'size': stat.st_size,
                    'modified': stat.st_mtime,
                    'page_count': len(reader.pages),
                    'source_type': 'file'
                }
            }
        except ImportError:
            logger.error("pypdf not installed, cannot load PDF files")
            raise ImportError("pypdf not installed, cannot load PDF files")
    
    def _load_docx(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Load a DOCX file.
        
        Args:
            file_path: Path to the DOCX file.
            **kwargs: Additional arguments.
        
        Returns:
            Dict[str, Any]: The loaded document.
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                text += "\n"
            
            # Get file metadata
            stat = os.stat(file_path)
            
            return {
                'text': text,
                'metadata': {
                    'source': file_path,
                    'filename': os.path.basename(file_path),
                    'extension': os.path.splitext(file_path)[1].lower(),
                    'size': stat.st_size,
                    'modified': stat.st_mtime,
                    'source_type': 'file'
                }
            }
        except ImportError:
            logger.error("python-docx not installed, cannot load DOCX files")
            raise ImportError("python-docx not installed, cannot load DOCX files")
    
    def _load_image(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Load an image file and extract text using OCR.
        
        Args:
            file_path: Path to the image file.
            **kwargs: Additional arguments.
        
        Returns:
            Dict[str, Any]: The loaded document.
        """
        try:
            import pytesseract
            from PIL import Image
            
            # Open image
            image = Image.open(file_path)
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            # Get file metadata
            stat = os.stat(file_path)
            
            return {
                'text': text,
                'metadata': {
                    'source': file_path,
                    'filename': os.path.basename(file_path),
                    'extension': os.path.splitext(file_path)[1].lower(),
                    'size': stat.st_size,
                    'modified': stat.st_mtime,
                    'image_size': f"{image.width}x{image.height}",
                    'source_type': 'file'
                }
            }
        except ImportError:
            logger.error("pytesseract or PIL not installed, cannot load image files")
            raise ImportError("pytesseract or PIL not installed, cannot load image files")
